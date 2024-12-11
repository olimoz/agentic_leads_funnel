import boto3
from botocore.exceptions import ClientError
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.application import MIMEApplication
import pandas as pd
from datetime import datetime
import os
from docx import Document
from docx.shared import Pt
from agents.results_ranking import SearchResultsRankingAgent

#======================================================================================
# EMAIL MANAGER FOR AMAZON SES EMAILS
#======================================================================================

class EmailManager:
    """
    Manages email-related operations including candidate selection for emails and email sending via Amazon SES.
    """
    def __init__(self, config_client, llm_email, client_folder, app):
        """
        Initialize the EmailManager.

        Args:
            config_client (dict): Client configuration containing email settings.
            app (WebResearchApp): The main application instance.
        """
        self.email_user = config_client['EMAIL_USER']
        self.email_recipient = config_client['EMAIL_BATCH_RECIPIENT']
        self.max_emails = config_client['MAX_EMAILS']
        self.client_folder = client_folder
        self.llm_email = llm_email
        self.app = app

    async def get_email_candidates(self, df_search_history, search_date_most_recent=None):
        """
        Retrieves email candidates from search history based on novelty score, search date, and other criteria.

        Args:
            df_search_history (pd.DataFrame): DataFrame containing search history.
            search_date_most_recent (datetime, optional): The most recent search date to use for NaT values.

        Returns:
            pd.DataFrame: DataFrame containing selected email candidates.
        """
        
        # Create a copy of the DataFrame, so we don't try updating slices, which is ambiguous.
        df = df_search_history.copy()
        
        # Ensure search_date is in datetime format
        df['search_date'] = pd.to_datetime(df['search_date'], errors='coerce')
        
        # Handle NaT values in search_date
        if search_date_most_recent is not None:
            search_date_most_recent = pd.to_datetime(search_date_most_recent)
            df.loc[df['search_date'].isna(), 'search_date'] = search_date_most_recent
        else:
            max_date = df['search_date'].max()
            df.loc[df['search_date'].isna(), 'search_date'] = max_date
        
        # First, create a condition to check if an individual's details were emailed recently
        def get_last_email_date(row, df):
            same_person = (
                (df['first_name'] == row['first_name']) & 
                (df['last_name'] == row['last_name']) & 
                (df['company'] == row['company']) &
                (df['email_date'].notna())  # only consider rows where an email was sent
            )
            if not df[same_person].empty:
                last_email = df[same_person]['email_date'].max()
                return (self.app.today - last_email).days > 31
            return True

        # no apply filters to get candidates
        filtered_df = df[
            (df['novelty_score'] > 5) & # reasonably novel
            (df['email_date'].isna()) & # the results of this search havn't been emailed
            ((self.app.today - df['search_date']).dt.days <= 31) & # search is less than a month old
            df.apply(lambda row: get_last_email_date(row, df), axis=1)  # no recent emails about this person
        ].copy()

        # Calculate the combined score
        filtered_df['combined_score'] = filtered_df['activity_score'] + filtered_df['services_need_score']
        
        # Sort by the combined score
        sorted_df = filtered_df.sort_values('combined_score', ascending=False)
        
        # Find the maximum score
        top_score = sorted_df['combined_score'].max()

        # Count how many rows have the top score
        top_score_count = sorted_df[sorted_df['combined_score'] == top_score].shape[0]

        # If all the top scores are the same, how do we differentiate?
        if top_score_count > self.max_emails:

            # we will rank the records, one batch at a time, using the SearchResultsRankingAgent
            batch_size = 6

            # create the ranking agent
            ranking_agent = SearchResultsRankingAgent(self.app)

            # blank list to receive rankings
            rankings = []
            consumption = []

            # for each batch of the data, up to the top_score_count
            for i in range(0, len(sorted_df[0:top_score_count]), batch_size):

                # get the batch
                batch = sorted_df.iloc[i:i+batch_size]

                # prepare texts for ranking
                texts = ""

                for _, row in batch.iterrows():
                    texts += f'<text first_name="{row["first_name"]}", last_name="{row["last_name"]}", company="{row["company"]}">\n'
                    texts += f'{row["search_results"]}\n'
                    texts += '</text>\n'
                texts += f'There are {top_score_count} texts. You must return {top_score_count} rankings'

                # get the ranking agent's results
                batch_rankings, consumption_item = await ranking_agent.run(self.llm_email, texts, self.client_folder)

                # add the rankings to the list
                rankings = rankings + batch_rankings
                consumption.append(consumption_item)

            # rankings is a list of dicts, so convert to pandas
            rank_df = pd.DataFrame(rankings)

            # Merge rankings with sorted_df on (first_name, last_name, company)
            merged_df = pd.merge(sorted_df, rank_df, on=['first_name', 'last_name', 'company'], how='left')

            # Sort the merged DataFrame by the rank column (NaN ranks will appear last)
            sorted_df = merged_df.sort_values(by=['combined_score', 'rank'], ascending=[False, True])

        else:
            # no consumption
            consumption = [{
                'function': 'get_email_candidates',
                'model': 'get_email_candidates',
                'search_calls': 0,
                'input_tokens': 0,
                'output_tokens': 0
            }]

        # Determine number of rows to select
        num_rows = min(int(len(sorted_df)), self.max_emails)
        
        # Select top candidates
        email_candidates = sorted_df.head(num_rows)
        
        # Select essential columns, keeping the original index
        email_candidates = email_candidates[['first_name', 'last_name', 'company', 'search_date', 'search_results', 'total_score']]

        # we need indices to update df_search_history, so get indices of email_candidates from df_search_history
        merged_df = df_search_history.merge(email_candidates, on=['first_name', 'last_name', 'company', 'search_date'], how='inner')
        mask = df_search_history.set_index(['first_name', 'last_name', 'company', 'search_date']).index.isin(
            email_candidates.set_index(['first_name', 'last_name', 'company', 'search_date']).index)
        email_candidates = df_search_history[mask]
        
        # log outcome
        self.app.logger.info(f"Email candidates selected, Qty={len(email_candidates)}")

        return email_candidates, consumption

    # function for use in creating final report
    def add_hyperlink(self, paragraph, url, text):

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id, )

        # Create a w:r element
        new_run = OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = OxmlElement('w:rPr')

        # Add color if it doesn't exist
        if not rPr.find(qn('w:color')):
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '0000FF')  # Blue color
            rPr.append(color)

        # Add underline if it doesn't exist
        if not rPr.find(qn('w:u')):
            u = OxmlElement('w:u')
            u.set(qn('w:val'), 'single')
            rPr.append(u)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

        return hyperlink

    # function for use in creating final report
    def add_hyperlinked_paragraph(self, doc, text):
        paragraph = doc.add_paragraph()
        
        # First, clean up the text by removing newlines around URLs in parentheses
        text = re.sub(r'\(\s*(\bhttps?://[^\s\)]+)\s*\)', r'(\1)', text)
        
        # Regular expression to find URLs, including those in parentheses
        url_pattern = r'https?://[^\s()[\]{},]+(?:\([^\s()[\]{},]*\))?[^\s()[\]{},]*(?:,(?=[^\s,]))?[^\s()[\]{},]*'
        
        # Find all URLs in the text
        urls = re.findall(url_pattern, text)
        
        # Split the text by URLs
        parts = re.split(url_pattern, text)
        
        # Add the text before the first URL
        paragraph.add_run(parts[0].strip())
        
        # Add a double line break before the first URL
        paragraph.add_run('\n\n')
        
        for i, url in enumerate(urls):
            # Strip any surrounding whitespace from the URL
            url = url.strip()
            
            # Add the URL as a hyperlink
            self.add_hyperlink(paragraph, url, url)
            
            # Add a line break after the URL
            paragraph.add_run('\n')
            
            # Add any text between this URL and the next one
            if i + 1 < len(parts):
                paragraph.add_run(parts[i + 1].strip())

    # Creating the final report for attachment to the email
    def create_word_report(self, df_updated_rows, attachment_path, searches_since_last_email):
        doc = Document()

        # Add title
        doc.add_heading('Proposed Leads', level=1)

        # Add date
        today = self.app.today.strftime("%a %d-%b-%Y")
        doc.add_paragraph(f'As of {today}')

        # Add introduction paragraph
        doc.add_paragraph(f'Since the last email, {searches_since_last_email} searches were completed. These are the most promising leads:')

        # Add Contents section
        doc.add_heading('Contents', level=2)
        for i, row in df_updated_rows.iterrows():
            doc.add_paragraph(f"{row['company']}, {row['first_name']} {row['last_name']}", style='List Number')

        # Add Results section
        doc.add_heading('Results', level=2)
        for _, row in df_updated_rows.iterrows():

            # Company and name
            doc.add_heading(f"{row['company']}, {row['first_name']} {row['last_name']}", level=3)
            
            # Search Results
            search_date = row['search_date'].strftime("%a %d-%b-%Y")
            # doc.add_paragraph().add_run('Search Results').bold = True
            doc.add_paragraph(f"Search Date: {search_date}")
            self.add_hyperlinked_paragraph(doc, row['search_results'].replace('\\n', '\n'))

            # Email subject
            subject = doc.add_paragraph()
            subject.add_run("Proposed Email").bold = True
            
            # Email content
            doc.add_paragraph(row['email_content'])
                    
            # Add horizontal line
            doc.add_paragraph('_' * 50)

        # End of Report
        doc.add_heading('End of Report', level=2)

        # Save the document to a BytesIO object
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Use the StorageManager to write the file
        today = self.app.today.strftime("%Y%m%d")
        self.app.storage_manager.write_file_binary(doc_bytes.getvalue(), attachment_path, self.client_folder)


    async def send_email(self, email_content, email_subject, attachment_path=None, client=None):
        """
        Sends an email with an optional attachment using AWS SES.
        Supports various file types including .docx, .xlsx, and .txt files.

        Args:
            email_content (str): The body of the email.
            email_subject (str): The subject of the email.
            attachment_path (str, optional): Path to the file to be attached.
            client (str, optional): The client name, used as a subfolder in StorageManager.

        Returns:
            str: A message indicating the result of the email sending operation.

        Raises:
            ValueError: If there's an error with the attachment or email sending.
        """
        ses_client = boto3.client('ses', region_name='eu-north-1')
        
        msg = MIMEMultipart('mixed')
        msg['Subject'] = email_subject
        msg['From'] = self.email_user
        msg['To'] = self.email_recipient
        
        # Attach the email body
        msg_body = MIMEMultipart('alternative')
        textpart = MIMEText(email_content, 'plain', 'utf-8')
        msg_body.attach(textpart)
        msg.attach(msg_body)
        
        if attachment_path is not None:
            try:
                attachment_content = self.app.storage_manager.read_file(attachment_path, client)
                if attachment_content is None:
                    msg = f"Failed to read attachment: {attachment_path}"
                    self.app.handle_error(self.app.logger, logging.ERROR, msg)
                
                filename = os.path.basename(attachment_path)
                
                # Determine the MIME type based on the file extension
                mime_type, _ = mimetypes.guess_type(filename)
                if mime_type is None:
                    mime_type = "application/octet-stream"
                
                # Ensure attachment_content is in bytes
                if isinstance(attachment_content, str):
                    attachment_content = attachment_content.encode('utf-8')
                elif isinstance(attachment_content, bytes):
                    pass
                else:
                    msg = f"Unexpected attachment content type: {type(attachment_content)}"
                    self.app.handle_error(self.app.logger, logging.ERROR, msg)
                
                # Handle different file types
                if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    # Word documents (.docx)
                    part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
                    part.set_payload(attachment_content)
                elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                    # Excel files (.xlsx)
                    part = MIMEBase('application', 'vnd.openxmlformats-officedocument.spreadsheetml.sheet')
                    part.set_payload(attachment_content)
                elif mime_type == 'text/plain':
                    # Text files
                    part = MIMEText(attachment_content.decode('utf-8'), 'plain', 'utf-8')
                else:
                    # Default handling for other file types
                    part = MIMEBase(*mime_type.split('/'))
                    part.set_payload(attachment_content)
                
                # Encode the payload if not already encoded
                if not isinstance(part, MIMEText):
                    encoders.encode_base64(part)
                
                # Log size after encoding
                encoded_size = len(part.get_payload().encode('utf-8') if isinstance(part, MIMEText) 
                                else part.get_payload().encode('ascii'))
                self.app.logger.info(f"Attachment '{filename}', from {attachment_path}, encoded size: {encoded_size} bytes")

                # Add headers
                part.add_header(
                    'Content-Disposition',
                    'attachment',
                    filename=filename
                )
                msg.attach(part)
                
            except Exception as e:
                msg = f"Error attaching file: {str(e)}"
                self.app.handle_error(self.app.logger, logging.CRITICAL, msg)
                raise
        
        try:
            response = await asyncio.to_thread(
                ses_client.send_raw_email,
                Source=self.email_user,
                Destinations=[self.email_recipient],
                RawMessage={'Data': msg.as_string()}
            )
        except ClientError as e:
            msg = f"Email client error: {str(e)}"
            self.app.handle_error(self.app.logger, logging.CRITICAL, msg)
            raise
        
        return f"Email has been sent to: {self.email_recipient}"